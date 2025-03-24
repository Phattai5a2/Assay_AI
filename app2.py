#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Mar 24 07:48:40 2025

@author: phattai
"""

import streamlit as st
import requests
import docx
import difflib
import os
from streamlit_quill import st_quill


# Sử dụng OpenRouter API miễn phí
API_URL = "https://openrouter.ai/api/v1/chat/completions"
API_KEY = "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"  # Thay bằng API key miễn phí từ OpenRouter

# Danh sách user giả lập
USERS = {
    "teacher": "1",
    "student": "1"
}

# Hàm kiểm tra đăng nhập
def login():
    st.session_state["logged_in"] = False
    user = st.text_input("Tên đăng nhập:")
    password = st.text_input("Mật khẩu:", type="password")
    if st.button("Đăng nhập"):
        if user in USERS and USERS[user] == password:
            st.session_state["logged_in"] = True
            st.session_state["user"] = user
            st.session_state["role"] = "teacher" if user == "teacher" else "student"
            st.session_state["user_folder"] = f"essays/{user}"
            os.makedirs(st.session_state["user_folder"], exist_ok=True)
            st.success(f"Xin chào, {user}!")
        else:
            st.error("Sai tài khoản hoặc mật khẩu!")

# Hàm đăng xuất
def logout():
    st.session_state.clear()
    st.rerun()

# Hàm lưu bài làm của sinh viên dưới dạng .docx
def save_essay(student_text, filename):
    user_folder = st.session_state.get("user_folder", "essays/unknown")
    os.makedirs(user_folder, exist_ok=True)
    filepath = os.path.join(user_folder, filename)
    
    doc = docx.Document()
    doc.add_paragraph(student_text)
    doc.save(filepath)
    
    return filepath

# Hàm đọc file Word
def read_docx(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# Hàm tính phần trăm giống nhau giữa hai văn bản
def calculate_similarity(text1, text2):
    similarity = difflib.SequenceMatcher(None, text1, text2).ratio()
    return round(similarity * 100, 2)

# Hàm chấm điểm bằng OpenRouter API
def grade_essay(student_text, answer_text):
    similarity = calculate_similarity(student_text, answer_text)
    prompt = f"""Bạn là giáo viên. Hãy chấm bài sau đây.
    \n\nĐáp án mẫu:\n{answer_text}
    \n\nBài làm của học sinh:\n{student_text}
    \n\nHãy đưa ra số điểm (thang 10) và nhận xét chi tiết."""
    
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "mistralai/mistral-small-3.1-24b-instruct:free",
        "messages": [{"role": "system", "content": "Bạn là một giáo viên chấm bài chuyên nghiệp."},
                      {"role": "user", "content": prompt}],
        "temperature": 0.7
    }
    
    response = requests.post(API_URL, headers=headers, json=payload)
    if response.status_code == 200:
        grading_result = response.json()["choices"][0]["message"]["content"]
        return grading_result, similarity
    else:
        return f"Lỗi: {response.status_code} - {response.json()}", None

# Giao diện chính
st.title("Hệ thống chấm bài tự luận AI")

if "logged_in" not in st.session_state or not st.session_state["logged_in"]:
    login()
else:
    st.write(f"Xin chào, {st.session_state['user']}!")
    if st.button("Đăng xuất"):
        logout()
    
    role = st.session_state.get("role", "student")
    
    if role == "student":
        uploaded_exam = st.file_uploader("Chọn đề thi tự luận từ giảng viên", type=["docx"])
        if uploaded_exam:
            exam_text = read_docx(uploaded_exam)
            questions = exam_text.split("\n")
            st.session_state["num_questions"] = len(questions)
            st.session_state["questions"] = questions
            
            st.subheader("Đề thi:")
            st.write(exam_text)
        
        if "num_questions" in st.session_state:
            answers = []
            for i, question in enumerate(st.session_state["questions"]):
                st.write(f"**{question}**")
                answer = st_quill(f"Câu {i+1}:", key=f"answer_{i}")
                answers.append(answer)
            
            if st.button("Nộp bài"):
                student_text = "\n".join(answers)
                filename = "essay.docx"
                filepath = save_essay(student_text, filename)
                st.success(f"Bài làm đã được lưu tại: {filepath}")
    
    elif role == "teacher":
        uploaded_exam = st.file_uploader("Tải lên đề thi tự luận", type=["docx"])
        if uploaded_exam:
            exam_path = os.path.join("essays", "exam.docx")
            with open(exam_path, "wb") as f:
                f.write(uploaded_exam.read())
            st.success("Đề thi đã được lưu.")
        
        student_folders = [f for f in os.listdir("essays") if os.path.isdir(os.path.join("essays", f))]
        selected_student = st.selectbox("Chọn sinh viên để chấm bài:", student_folders)
        if selected_student:
            essay_path = os.path.join("essays", selected_student, "essay.docx")
            if os.path.exists(essay_path):
                student_text = read_docx(essay_path)
                
                uploaded_answer = st.file_uploader("Tải lên file đáp án mẫu (Word)", type=["docx"])
                if uploaded_answer:
                    answer_text = read_docx(uploaded_answer)
                    if st.button("Chấm điểm"):
                        with st.spinner("Đang chấm điểm..."):
                            result, similarity = grade_essay(student_text, answer_text)
                        st.subheader("Kết quả chấm bài")
                        st.write(result)
                        if similarity is not None:
                            st.write(f"Mức độ giống đáp án: {similarity}%")