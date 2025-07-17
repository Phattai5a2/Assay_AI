import streamlit as st
import requests
import docx
import pandas as pd
from streamlit_quill import st_quill
import base64
import io
import json
import re
import zipfile
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request

# Khởi tạo các biến trạng thái
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False
if "user" not in st.session_state:
    st.session_state["user"] = None
if "role" not in st.session_state:
    st.session_state["role"] = None
if "uploaded_files" not in st.session_state:
    st.session_state["uploaded_files"] = []
if "grading_results" not in st.session_state:
    st.session_state["grading_results"] = []
if "start_exam" not in st.session_state:
    st.session_state["start_exam"] = False
if "current_num_questions" not in st.session_state:
    st.session_state["current_num_questions"] = 1
if "mssv" not in st.session_state:
    st.session_state["mssv"] = ""
if "full_name" not in st.session_state:
    st.session_state["full_name"] = ""
if "exam_access_granted" not in st.session_state:
    st.session_state["exam_access_granted"] = False
if "upload_completed" not in st.session_state:
    st.session_state["upload_completed"] = False

st.markdown(
    """
    <style>
    [data-testid="stToolbar"] {
            visibility: hidden;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Sử dụng OpenRouter API miễn phí
API_URL = "https://openrouter.ai/api/v1/chat/completions"
try:
    API_KEY = st.secrets["openrouter"]["api_key"]
except KeyError:
    st.error("Không tìm thấy API key của OpenRouter trong Secrets. Vui lòng thêm 'openrouter.api_key' vào Secrets trên Streamlit Cloud.")
    st.stop()

# Hàm thay đổi con trỏ chuột
def set_loading_cursor(status):
    if status:
        st.markdown(
            """
            <style>
            html, body {
                cursor: wait !important;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            """
            <style>
            html, body {
                cursor: default !important;
            }
            </style>
            """,
            unsafe_allow_html=True
        )

# Hàm loại bỏ các ký tự ### và #### từ nội dung Markdown
def clean_markdown_headers(text):
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        line = line.replace("### ", "").replace("#### ", "")
        line = line.replace("** ", "").replace("**", "")
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

# Xác thực Google Drive
def authenticate_google_drive():
    SCOPES = ['https://www.googleapis.com/auth/drive']
    
    try:
        creds_info_str = st.secrets["google_drive"]["credentials"]
        client_secrets_str = st.secrets["google_drive"]["client_secrets"]
        creds_info = json.loads(creds_info_str)
        client_secrets = json.loads(client_secrets_str)
    except KeyError:
        error_msg = (
            "Không tìm thấy thông tin xác thực trong Secrets.\n"
            "Vui lòng thêm client_secrets và credentials vào Secrets trên Streamlit Cloud."
        )
        st.error(error_msg)
        st.stop()
    except json.JSONDecodeError as e:
        error_msg = (
            "Dữ liệu trong Secrets không đúng định dạng JSON.\n"
            f"Chi tiết lỗi: {str(e)}\n"
            "Vui lòng kiểm tra lại client_secrets và credentials trong Secrets trên Streamlit Cloud."
        )
        st.error(error_msg)
        st.stop()

    creds = None
    try:
        creds = Credentials.from_authorized_user_info(info=creds_info, scopes=SCOPES)
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception as e:
                st.error(f"Lỗi khi làm mới token: {str(e)}")
                st.error("Vui lòng cập nhật credentials mới trong Secrets trên Streamlit Cloud.")
                st.stop()
    except Exception as e:
        st.error(f"Lỗi khi tạo credentials: {str(e)}")
        st.error("Vui lòng kiểm tra hoặc cập nhật credentials trong Secrets trên Streamlit Cloud.")
        st.stop()
    
    if not creds or not creds.valid:
        st.error("Credentials không hợp lệ. Vui lòng cập nhật credentials mới trong Secrets trên Streamlit Cloud.")
        st.stop()
    
    service = build('drive', 'v3', credentials=creds)
    return service

# Tạo hoặc lấy ID của thư mục trên Google Drive
def get_or_create_folder(service, folder_name, parent_id=None):
    query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    response = service.files().list(q=query, spaces='drive').execute()
    folders = response.get('files', [])
    
    if folders:
        return folders[0]['id']
    else:
        folder_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_id:
            folder_metadata['parents'] = [parent_id]
        folder = service.files().create(body=folder_metadata, fields='id').execute()
        return folder['id']

# Xóa tất cả file trong một thư mục trên Google Drive
def clear_folder(service, folder_id):
    try:
        response = service.files().list(q=f"'{folder_id}' in parents and trashed=false", spaces='drive').execute()
        file_list = response.get('files', [])
        for file in file_list:
            service.files().delete(fileId=file['id']).execute()
    except Exception as e:
        st.error(f"Lỗi khi xóa file trong thư mục: {str(e)}")

# Tải file lên Google Drive và đặt quyền chia sẻ công khai
def upload_file_to_drive(service, file_content, file_name, folder_id, update_if_exists=True):
    # Tìm tất cả các file có tên bắt đầu bằng {mssv}_{student_name}_graded
    try:
        # Trích xuất MSSV và tên sinh viên từ tên file
        base_name = file_name.replace("_graded.docx", "")
        query = f"'{folder_id}' in parents and trashed=false"
        response = service.files().list(q=query, spaces='drive').execute()
        files = response.get('files', [])
        
        # Xóa tất cả các file có tên bắt đầu bằng base_name
        for file in files:
            if file['name'].startswith(base_name):
                service.files().delete(fileId=file['id']).execute()
    except Exception as e:
        st.error(f"Lỗi khi xóa file trùng tên: {str(e)}")
        return None
    
    # Tải file mới lên
    file_metadata = {
        'name': file_name,
        'parents': [folder_id]
    }
    media = MediaIoBaseUpload(io.BytesIO(file_content), mimetype='application/octet-stream')
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    
    file_id = file['id']
    try:
        permission = {
            'type': 'anyone',
            'role': 'reader'
        }
        service.permissions().create(fileId=file_id, body=permission).execute()
    except Exception as e:
        st.error(f"Không thể đặt quyền chia sẻ công khai cho file {file_name}: {str(e)}")
        raise
    
    return file_id

# Tải file từ Google Drive
def download_file_from_drive(service, file_id):
    try:
        request = service.files().get_media(fileId=file_id)
        file_content = io.BytesIO()
        downloader = MediaIoBaseDownload(file_content, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        file_content.seek(0)
        return file_content.read()
    except Exception as e:
        st.error(f"Lỗi khi tải file từ Google Drive: {str(e)}")
        return None

# Tìm file trên Google Drive
def find_file_in_folder(service, file_name, folder_id):
    try:
        query = f"name='{file_name}' and '{folder_id}' in parents and trashed=false"
        response = service.files().list(q=query, spaces='drive').execute()
        files = response.get('files', [])
        return files[0] if files else None
    except Exception as e:
        st.error(f"Lỗi khi tìm file trên Google Drive: {str(e)}")
        return None

# Lấy danh sách user từ file users.json
def load_users(service, root_folder_id):
    try:
        users_file = find_file_in_folder(service, "users.json", root_folder_id)
        if users_file:
            content = download_file_from_drive(service, users_file['id'])
            if content:
                return json.loads(content.decode('utf-8'))
            else:
                st.error("Không thể đọc nội dung file users.json.")
                return []
        else:
            # Nếu chưa có file, tạo file với user admin mặc định
            default_users = [
                {"username": "admin", "password": "admin123", "role": "admin"},
                {"username": "teacher", "password": "1", "role": "teacher"},
                {"username": "student", "password": "1", "role": "student"},
                {"username": "teacher2", "password": "1", "role": "teacher"},
                {"username": "tai", "password": "1", "role": "teacher"}
            ]
            save_users(service, root_folder_id, default_users)
            st.info("Đã tạo file users.json với user admin mặc định (admin/admin123).")
            return default_users
    except Exception as e:
        st.error(f"Lỗi khi tải danh sách user: {str(e)}")
        return []

# Lưu danh sách user vào file users.json
def save_users(service, root_folder_id, users):
    try:
        json_content = json.dumps(users, ensure_ascii=False, indent=4)
        upload_file_to_drive(service, json_content.encode('utf-8'), "users.json", root_folder_id, update_if_exists=True)
    except Exception as e:
        st.error(f"Lỗi khi lưu danh sách user: {str(e)}")

# Lấy danh sách đề thi từ thư mục của giảng viên
def get_exam_list(service, exams_folder_id):
    try:
        exam_secrets_file = find_file_in_folder(service, "exam_secrets.json", exams_folder_id)
        if exam_secrets_file:
            content = download_file_from_drive(service, exam_secrets_file['id'])
            if content:
                return json.loads(content.decode('utf-8'))
            else:
                st.error("Không thể đọc nội dung file exam_secrets.json.")
                return []
        return []
    except Exception as e:
        st.error(f"Lỗi khi tải danh sách đề thi: {str(e)}")
        return []

# Cập nhật danh sách đề thi vào file exam_secrets.json
def update_exam_list(service, exams_folder_id, exam_list):
    try:
        json_content = json.dumps(exam_list, ensure_ascii=False, indent=4)
        upload_file_to_drive(service, json_content.encode('utf-8'), "exam_secrets.json", exams_folder_id, update_if_exists=True)
    except Exception as e:
        st.error(f"Lỗi khi lưu danh sách đề thi: {str(e)}")

# Khởi tạo Google Drive
try:
    service = authenticate_google_drive()
except Exception as e:
    st.error(f"Lỗi khi khởi tạo Google Drive: {str(e)}")
    st.stop()

# Tạo thư mục gốc
root_folder_id = get_or_create_folder(service, "ExamSystem")
if not root_folder_id:
    st.error("Không thể tạo hoặc truy cập thư mục ExamSystem trên Google Drive.")
    st.stop()

# Tạo thư mục riêng cho từng giảng viên
def initialize_teacher_folders(service, username):
    teacher_folder = get_or_create_folder(service, f"teacher_{username}", root_folder_id)
    exams_folder = get_or_create_folder(service, "exams", teacher_folder)
    essays_folder = get_or_create_folder(service, "essays", teacher_folder)
    graded_essays_folder = get_or_create_folder(service, "graded_essays", teacher_folder)
    reports_folder = get_or_create_folder(service, "reports", teacher_folder)
    return {
        "teacher_folder_id": teacher_folder,
        "exams_folder_id": exams_folder,
        "essays_folder_id": essays_folder,
        "graded_essays_folder_id": graded_essays_folder,
        "reports_folder_id": reports_folder
    }

# Hàm kiểm tra đăng nhập
def login():
    st.session_state["logged_in"] = False
    st.markdown(
        """
        <h2 style='text-align: center; font-size: 36px;'>👤Đăng nhập hệ thống</h2>
        """,
        unsafe_allow_html=True
    )
    user = st.text_input("Tên đăng nhập:")
    password = st.text_input("Mật khẩu:", type="password")
    if st.button("Đăng nhập", icon=":material/login:"):
        if not user or not password:
            st.error("Vui lòng nhập đầy đủ tên đăng nhập và mật khẩu.")
            return
        
        users = load_users(service, root_folder_id)
        if not users:
            st.error("Không thể tải danh sách user. Vui lòng kiểm tra kết nối Google Drive.")
            return
        
        user_data = next((u for u in users if u["username"] == user and u["password"] == password), None)
        if user_data:
            st.session_state["logged_in"] = True
            st.session_state["user"] = user
            st.session_state["role"] = user_data["role"]
            st.success(f"Xin chào, {user}!")
            st.rerun()
        else:
            st.error("Sai tài khoản hoặc mật khẩu! Vui lòng kiểm tra lại.")

# Hàm đăng xuất
def logout():
    st.session_state.clear()
    st.rerun()

# Hàm đọc file Word
def read_docx(file_content):
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Lỗi khi đọc file Word: {str(e)}")
        return ""

# Hàm lưu vào CSV trên Google Drive với mã hóa UTF-8-SIG
def save_to_csv(data, service, folder_id):
    try:
        df = pd.DataFrame(data)
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
        
        existing_file = find_file_in_folder(service, "grading_report.csv", folder_id)
        if existing_file:
            existing_content = download_file_from_drive(service, existing_file['id'])
            if existing_content:
                existing_df = pd.read_csv(io.StringIO(existing_content.decode('utf-8-sig')), encoding='utf-8-sig')
                df = pd.concat([existing_df, df], ignore_index=True)
                df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
                file_metadata = {'name': "grading_report.csv"}
                media = MediaIoBaseUpload(io.BytesIO(csv_buffer.getvalue().encode('utf-8')), mimetype='text/csv')
                service.files().update(fileId=existing_file['id'], body=file_metadata, media_body=media).execute()
        else:
            df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
            upload_file_to_drive(service, csv_buffer.getvalue().encode('utf-8'), "grading_report.csv", folder_id)
    except Exception as e:
        st.error(f"Lỗi khi lưu file CSV: {str(e)}")

# Hàm chấm điểm bài tự luận
def grade_essay(student_text, answer_text, student_name=None, mssv=None):
    prompt = f"""Bạn là một giảng viên chấm bài chuyên nghiệp. Hãy chấm bài tự luận sau đây.

    **Đáp án mẫu:**
    {answer_text}

    **Bài làm của sinh viên:**
    {student_text}

    **Yêu cầu chấm bài:**
    1. Đưa ra nhận xét chi tiết về bài làm của sinh viên, bao gồm nhận xét cho từng câu (nếu có).
    2. Chấm điểm trên thang điểm 10 (Điểm tổng của sinh viên luôn không quá 10 điểm), điểm từng câu phải nhỏ hơn hay bằng điểm trong từng câu ghi trong đáp án với định dạng: **Điểm: [số điểm]** (ví dụ: Điểm: 5.0).
    3. Cuối cùng, ghi rõ tổng điểm của bài làm theo định dạng: **Tổng điểm: [số điểm]** (ví dụ: Tổng điểm: 6.0). 
       - Dòng này phải là dòng cuối cùng.
       - Không thêm bất kỳ từ ngữ nào khác trước hoặc sau (ví dụ: không ghi "Tổng điểm ghi là", "Kết luận", v.v.).

    **Ví dụ định dạng kết quả:**
    Nhận xét chi tiết về bài làm của sinh viên:

    **Câu 1:**
    - Sinh viên giải thích đúng khái niệm.
    - Điểm trừ: Thiếu ví dụ bổ sung.
    
    **Câu 2:**
    - Sinh viên mô tả đúng một phần.
    - Điểm trừ: Thiếu giải thích chi tiết.

    Điểm:
    - Câu 1: **3.0**
    - Câu 2: **2.5**

    **Tổng điểm: 6.5**

    Bắt đầu chấm bài:"""
    
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "mistralai/mistral-small-3.1-24b-instruct:free",
        "messages": [{"role": "system", "content": "Bạn là một giảng viên chấm bài chuyên nghiệp."},
                     {"role": "user", "content": prompt}],
        "temperature": 0.3  # Giảm temperature để AI tuân thủ prompt chặt chẽ hơn
    }
    
    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=30)
        if response.status_code == 200:
            grading_result = response.json()["choices"][0]["message"]["content"]
            if student_name and mssv:
                total_score = extract_score(grading_result)
                data = {
                    "MSSV": [mssv],
                    "Họ và Tên": [student_name],
                    "Điểm Tổng": [total_score],
                    "Kết quả chấm điểm": [grading_result]
                }
                save_to_csv(data, service, reports_folder_id)
            return grading_result
        else:
            error_detail = response.json() if response.content else "No response content"
            st.error(f"Lỗi API: {response.status_code} - {error_detail}")
            return None
    except requests.exceptions.Timeout:
        st.error("Yêu cầu API đã hết thời gian (timeout). Vui lòng thử lại sau.")
        return None
    except requests.exceptions.RequestException as e:
        st.error(f"Lỗi kết nối mạng: {str(e)}")
        return None

# Hàm trích xuất điểm từ kết quả chấm
def extract_score(grading_result):
    # Tìm tất cả các lần xuất hiện của "Tổng điểm:" và lấy lần cuối cùng
    matches = re.findall(r"Tổng điểm:\s*(\d+(\.\d+)?)", grading_result, re.IGNORECASE)
    if matches:
        return float(matches[-1][0])  # Lấy điểm số từ lần xuất hiện cuối cùng
    
    # Nếu không tìm thấy "Tổng điểm:", tìm "Điểm:"
    matches = re.findall(r"Điểm:\s*(\d+(\.\d+)?)", grading_result, re.IGNORECASE)
    if matches:
        return float(matches[-1][0])  # Lấy điểm số từ lần xuất hiện cuối cùng của "Điểm:"
    
    # Tìm định dạng: Điểm: 5.5/9 (trích xuất 5.5)
    matches = re.findall(r"Điểm:\s*(\d+(\.\d+)?)/\d+", grading_result, re.IGNORECASE)
    if matches:
        return float(matches[-1][0])
    
    # Tìm định dạng: Score: 5.5
    matches = re.findall(r"Score:\s*(\d+(\.\d+)?)", grading_result, re.IGNORECASE)
    if matches:
        return float(matches[-1][0])
    
    # Tìm định dạng: 5.5/10
    matches = re.findall(r"(\d+(\.\d+)?)/10", grading_result)
    if matches:
        return float(matches[-1][0])
    
    # Tìm định dạng: Một dòng chỉ chứa số (ví dụ: 5.5)
    matches = re.findall(r"^\s*(\d+(\.\d+)?)\s*$", grading_result, re.MULTILINE)
    if matches:
        return float(matches[-1][0])
    
    st.warning(f"Không thể trích xuất điểm từ kết quả: {grading_result}")
    return 0.0

# Hàm đọc báo cáo từ Google Drive với mã hóa UTF-8-SIG
def load_grading_report(service, folder_id):
    try:
        file = find_file_in_folder(service, "grading_report.csv", folder_id)
        if file:
            content = download_file_from_drive(service, file['id'])
            if content:
                return pd.read_csv(io.StringIO(content.decode('utf-8-sig')), encoding='utf-8-sig')
        return None
    except Exception as e:
        st.error(f"Lỗi khi đọc báo cáo: {str(e)}")
        return None

# Giao diện chính
if not st.session_state["logged_in"]:
    login()
else:
    # Hiển thị tiêu đề dựa trên vai trò
    role = st.session_state.get("role", "student")
    if role == "student":
        st.markdown(
            "<h1 style='text-align: center; font-size: 40px;'>Hệ thống thi tự luận trực tuyến NTTU</h1>",
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            "<h1 style='text-align: center; font-size: 40px;'>🎓Hệ thống chấm tự luận bằng AI</h1>",
            unsafe_allow_html=True
        )
    
    st.write(f"Xin chào, {st.session_state['user']}!")
    if st.button("Đăng xuất"):
        logout()
    
    if role == "admin":
        st.subheader("Quản lý user")
        
        # Hiển thị danh sách user hiện có dưới dạng bảng
        users = load_users(service, root_folder_id)
        if users:
            st.info("Danh sách user hiện có:")
            
            # Tạo DataFrame từ danh sách user
            user_data = {
                "Tên đăng nhập": [user["username"] for user in users],
                "Vai trò": [user["role"] for user in users]
            }
            df = pd.DataFrame(user_data)
            
            # Thêm CSS để làm đẹp bảng
            st.markdown(
                """
                <style>
                .dataframe {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    font-size: 16px;
                    text-align: left;
                }
                .dataframe th {
                    background-color: #4CAF50;
                    color: white;
                    padding: 12px 15px;
                    text-align: center;
                    border: 1px solid #ddd;
                }
                .dataframe td {
                    padding: 12px 15px;
                    border: 1px solid #ddd;
                }
                .dataframe tr:nth-child(even) {
                    background-color: #f2f2f2;
                }
                .dataframe tr:hover {
                    background-color: #ddd;
                }
                </style>
                """,
                unsafe_allow_html=True
            )
            
            # Hiển thị bảng
            st.dataframe(df, use_container_width=True)
        else:
            st.error("Không thể tải danh sách user.")
        
        # Form đăng ký user mới
        st.subheader("Đăng ký user mới")
        new_username = st.text_input("Tên đăng nhập mới:")
        new_password = st.text_input("Mật khẩu mới:", type="password")
        new_role = st.selectbox("Vai trò:", ["admin", "teacher", "student"])
        
        if st.button("Đăng ký"):
            if not new_username or not new_password:
                st.error("Vui lòng nhập đầy đủ tên đăng nhập và mật khẩu.")
            else:
                # Kiểm tra username đã tồn tại chưa
                if any(user["username"] == new_username for user in users):
                    st.error("Tên đăng nhập đã tồn tại. Vui lòng chọn tên khác.")
                else:
                    # Thêm user mới
                    users.append({
                        "username": new_username,
                        "password": new_password,
                        "role": new_role
                    })
                    save_users(service, root_folder_id, users)
                    st.success(f"Đã đăng ký user {new_username} với vai trò {new_role}.")
                    st.rerun()
    
    elif role == "teacher":
        teacher_folders = initialize_teacher_folders(service, st.session_state["user"])
        exams_folder_id = teacher_folders["exams_folder_id"]
        essays_folder_id = teacher_folders["essays_folder_id"]
        graded_essays_folder_id = teacher_folders["graded_essays_folder_id"]
        reports_folder_id = teacher_folders["reports_folder_id"]

        # Tạo 2 tab cho giảng viên
        tab1, tab2 = st.tabs(["Tải đề thi lên", "Chấm bài thi tự luận"])

        # Tab 1: Tải đề thi lên
        with tab1:
            st.subheader("Tải đề thi và đáp án")

            # Hiển thị danh sách đề thi hiện có
            exam_list = get_exam_list(service, exams_folder_id)
            if exam_list:
                st.info("Danh sách đề thi hiện có:")
                for exam in exam_list:
                    subject_code = exam.get("subject_code", "N/A")
                    term = exam.get("term", "N/A")
                    subject_name = exam.get("subject_name", "N/A")
                    st.write(f"- {subject_code} - {term} - {subject_name} - {exam['exam_file']} (Mã số bí mật: {exam['secret_code']})")

            # Nút xóa tất cả đề thi
            col1, col2 = st.columns(2)
            with col1:
                if exam_list and st.button("Xóa tất cả đề thi"):
                    for exam in exam_list:
                        service.files().delete(fileId=exam['exam_id']).execute()
                        service.files().delete(fileId=exam['answer_id']).execute()
                    exam_secrets_file = find_file_in_folder(service, "exam_secrets.json", exams_folder_id)
                    if exam_secrets_file:
                        service.files().delete(fileId=exam_secrets_file['id']).execute()
                    st.success("Đã xóa tất cả đề thi và đáp án.")
                    st.rerun()

            # Form tải lên đề thi mới
            st.subheader("Tải lên đề thi mới")
            uploaded_exam_pdf = st.file_uploader("Tải lên đề thi (PDF)", type=["pdf"], key="exam_pdf")
            uploaded_answer = st.file_uploader("Tải lên đáp án mẫu", type=["docx"], key="answer")
            subject_code = st.text_input("Mã học phần (ví dụ: 012407662501):", key="subject_code")
            term = st.text_input("Tên lớn (ví dụ: 25DHT1A):", key="term")
            subject_name = st.text_input("Tên môn học (ví dụ: Nhập môn KHDL):", key="subject_name")
            secret_code = st.text_input("Nhập mã số bí mật cho đề thi:", type="password", key="secret_code")

            if st.button("Tải lên đề thi"):
                if not uploaded_exam_pdf or not uploaded_answer:
                    st.error("Vui lòng tải lên cả file đề thi (PDF) và đáp án mẫu (DOCX).")
                elif not subject_code or not term or not subject_name or not secret_code:
                    st.error("Vui lòng nhập đầy đủ Mã học phần, Tên lớn, Tên môn học và Mã số bí mật.")
                else:
                    exam_pdf_content = uploaded_exam_pdf.read()
                    answer_content = uploaded_answer.read()

                    exam_count = len(exam_list) + 1
                    exam_filename = f"de_thi_{exam_count}.pdf"
                    answer_filename = f"dap_an_{exam_count}.docx"

                    exam_file_id = upload_file_to_drive(service, exam_pdf_content, exam_filename, exams_folder_id, update_if_exists=True)
                    answer_file_id = upload_file_to_drive(service, answer_content, answer_filename, exams_folder_id, update_if_exists=True)

                    exam_list.append({
                        "exam_file": exam_filename,
                        "exam_id": exam_file_id,
                        "answer_file": answer_filename,
                        "answer_id": answer_file_id,
                        "secret_code": secret_code,
                        "subject_code": subject_code,
                        "term": term,
                        "subject_name": subject_name
                    })
                    update_exam_list(service, exams_folder_id, exam_list)

                    st.success(f"Đề thi {exam_filename} và đáp án đã được lưu trên Google Drive.")
                    st.rerun()

        # Tab 2: Chấm bài thi tự luận
        with tab2:
            st.subheader("Chấm bài thi tự luận bằng AI")

            # Hiển thị danh sách đề thi để chọn
            exam_list = get_exam_list(service, exams_folder_id)
            if not exam_list:
                st.error("Không tìm thấy đề thi nào trên Google Drive. Vui lòng tải lên đề thi và đáp án mẫu trong tab 'Tải đề thi lên'.")
            else:
                st.info("Danh sách đề thi hiện có:")
                display_names = [f"{exam['subject_code']} - {exam['term']} - {exam['subject_name']}" for exam in exam_list]
                selected_display_name = st.selectbox("Chọn đề thi và đáp án mẫu:", display_names, key="select_exam_tab2")

                # Tìm exam tương ứng với display_name đã chọn
                selected_exam = next(exam for exam in exam_list 
                                   if f"{exam['subject_code']} - {exam['term']} - {exam['subject_name']}" == selected_display_name)
                answer_content = download_file_from_drive(service, selected_exam['answer_id'])
                answer_text = read_docx(answer_content)

                # Lấy danh sách bài làm từ thư mục essays
                response = service.files().list(q=f"'{essays_folder_id}' in parents and trashed=false", spaces='drive').execute()
                file_list = response.get('files', [])
                
                essay_data = []
                if file_list:
                    for file in file_list:
                        if file['name'].endswith(".docx"):
                            try:
                                mssv, student_name = file['name'].replace(".docx", "").split("_", 1)
                                essay_data.append({
                                    "MSSV": mssv,
                                    "Họ và Tên": student_name,
                                    "Tên file": file['name'],
                                    "File ID": file['id']
                                })
                            except ValueError:
                                st.warning(f"Tên file {file['name']} không đúng định dạng 'MSSV_HọTên.docx'. Bỏ qua.")

                # Tạo 3 sub-tab: Chấm bài đơn, Chấm bài hàng loạt, Xem báo cáo
                sub_tab1, sub_tab2, sub_tab3 = st.tabs(["Chấm bài đơn", "Chấm bài hàng loạt", "Xem báo cáo"])

                # Sub-tab 1: Chấm bài đơn
                with sub_tab1:
                    if essay_data:
                        st.subheader("Chọn bài làm từ danh sách")
                        selected_essay = st.selectbox("Chọn bài làm để chấm:", [f"{data['MSSV']} - {data['Họ và Tên']}" for data in essay_data], key="select_single_essay_tab2")
                        selected_essay_data = next(data for data in essay_data if f"{data['MSSV']} - {data['Họ và Tên']}" == selected_essay)

                        # Nút tải bài làm về máy
                        if st.button("Tải bài làm này", key="download_single_essay_tab2"):
                            set_loading_cursor(True)
                            with st.spinner(f"Đang tải file {selected_essay_data['Tên file']}..."):
                                file_content = download_file_from_drive(service, selected_essay_data['File ID'])
                            set_loading_cursor(False)
                            if file_content:
                                st.download_button(
                                    label=f"Tải {selected_essay_data['Tên file']}",
                                    data=file_content,
                                    file_name=selected_essay_data['Tên file'],
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_single_essay_button_tab2"
                                )

                        # Nút chấm bài
                        if st.button("Chấm bài", key="grade_single_essay_tab2"):
                            set_loading_cursor(True)
                            with st.spinner(f"Đang tải bài làm {selected_essay_data['Tên file']}..."):
                                student_content = download_file_from_drive(service, selected_essay_data['File ID'])
                            set_loading_cursor(False)

                            if student_content:
                                student_text = read_docx(student_content)
                                mssv = selected_essay_data['MSSV']
                                student_name = selected_essay_data['Họ và Tên']

                                set_loading_cursor(True)
                                with st.spinner("Đang chấm bài..."):
                                    result = grade_essay(student_text, answer_text, student_name, mssv)
                                set_loading_cursor(False)

                                if result:
                                    st.subheader("Kết quả chấm điểm:")
                                    st.write(f"MSSV: {mssv}")
                                    st.write(f"Họ và Tên: {student_name}")
                                    st.write(result)

                                    # Loại bỏ các ký tự ### và #### trước khi lưu vào file Word
                                    cleaned_result = clean_markdown_headers(result)

                                    graded_filename = f"{mssv}_{student_name}_graded.docx"
                                    doc = docx.Document()
                                    doc.add_paragraph(f"MSSV: {mssv}")
                                    doc.add_paragraph(f"Họ và Tên: {student_name}")
                                    doc.add_paragraph(cleaned_result)
                                    doc_buffer = io.BytesIO()
                                    doc.save(doc_buffer)
                                    doc_buffer.seek(0)

                                    upload_file_to_drive(service, doc_buffer.getvalue(), graded_filename, graded_essays_folder_id)

                                    st.success(f"Kết quả đã được lưu trên Google Drive với tên: {graded_filename}")
                                    st.download_button(
                                        label="Tải kết quả chấm điểm",
                                        data=doc_buffer.getvalue(),
                                        file_name=graded_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_graded_single_essay_tab2"
                                    )
                    else:
                        st.info("Chưa có bài làm nào được nộp.")

                # Sub-tab 2: Chấm bài hàng loạt (đã sửa để đảm bảo chọn được nhiều bài)
                with sub_tab2:
                    if essay_data:
                        st.subheader("Chọn bài làm để chấm hàng loạt")
                        # Thêm thông báo kiểm tra số lượng bài làm
                        st.info(f"Số lượng bài làm hiện có: {len(essay_data)}")
                        if len(essay_data) < 2:
                            st.warning("Hiện tại chỉ có 1 bài làm hoặc không có bài làm nào. Vui lòng kiểm tra thư mục 'essays' trên Google Drive và đảm bảo các file có định dạng 'MSSV_HọTên.docx'.")

                        # Sử dụng st.multiselect để chọn nhiều bài làm
                        selected_essays = st.multiselect(
                            "Chọn các bài làm để chấm (có thể chọn nhiều bài):",
                            [f"{data['MSSV']} - {data['Họ và Tên']}" for data in essay_data],
                            key="select_batch_essays_tab2"
                        )

                        # Nút tải tất cả bài làm đã chọn
                        if selected_essays:
                            if st.button("Tải các bài làm đã chọn (ZIP)", key="download_batch_essays_tab2"):
                                set_loading_cursor(True)
                                with st.spinner("Đang tạo file ZIP..."):
                                    zip_buffer = io.BytesIO()
                                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                        for essay in selected_essays:
                                            essay_data_selected = next(data for data in essay_data if f"{data['MSSV']} - {data['Họ và Tên']}" == essay)
                                            file_content = download_file_from_drive(service, essay_data_selected['File ID'])
                                            if file_content:
                                                zip_file.writestr(essay_data_selected['Tên file'], file_content)
                                    zip_buffer.seek(0)
                                set_loading_cursor(False)
                                st.download_button(
                                    label="Tải các bài làm đã chọn (ZIP)",
                                    data=zip_buffer,
                                    file_name="selected_student_essays.zip",
                                    mime="application/zip",
                                    key="download_batch_essays_zip_tab2"
                                )

                        # Khởi tạo biến trạng thái cho việc chấm bài
                        if "start_grading" not in st.session_state:
                            st.session_state["start_grading"] = False

                        if selected_essays:
                            # Nút "Chấm bài" để bắt đầu quá trình chấm
                            if st.button("Chấm bài", key="grade_batch_essays_tab2"):
                                st.session_state["start_grading"] = True
                                st.session_state["grading_results"] = []  # Reset kết quả trước khi chấm

                                # Xóa tất cả file cũ trong thư mục graded_essays trước khi chấm
                                set_loading_cursor(True)
                                with st.spinner("Đang xóa các file kết quả cũ..."):
                                    clear_folder(service, graded_essays_folder_id)
                                set_loading_cursor(False)

                                results = []

                                set_loading_cursor(True)
                                with st.spinner("Đang chấm bài hàng loạt..."):
                                    for idx, essay in enumerate(selected_essays, 1):
                                        essay_data_selected = next(data for data in essay_data if f"{data['MSSV']} - {data['Họ và Tên']}" == essay)
                                        file_content = download_file_from_drive(service, essay_data_selected['File ID'])
                                        if file_content:
                                            student_text = read_docx(file_content)
                                            mssv = essay_data_selected['MSSV']
                                            student_name = essay_data_selected['Họ và Tên']

                                            grading_result = grade_essay(student_text, answer_text, student_name, mssv)

                                            if grading_result:
                                                total_score = extract_score(grading_result)
                                                results.append({
                                                    "STT": idx,
                                                    "MSSV": mssv,
                                                    "Họ và Tên": student_name,
                                                    "Tổng điểm tự luận": total_score
                                                })

                                                # Loại bỏ các ký tự ### và #### trước khi lưu vào file Word
                                                cleaned_result = clean_markdown_headers(grading_result)

                                                graded_filename = f"{mssv}_{student_name}_graded.docx"
                                                doc = docx.Document()
                                                doc.add_paragraph(f"MSSV: {mssv}")
                                                doc.add_paragraph(f"Họ và Tên: {student_name}")
                                                doc.add_paragraph(cleaned_result)
                                                doc_buffer = io.BytesIO()
                                                doc.save(doc_buffer)
                                                doc_buffer.seek(0)

                                                upload_file_to_drive(service, doc_buffer.getvalue(), graded_filename, graded_essays_folder_id)

                                set_loading_cursor(False)
                                st.session_state["grading_results"] = results

                        if st.session_state["grading_results"]:
                            df = pd.DataFrame(st.session_state["grading_results"])
                            st.subheader("Kết quả chấm điểm hàng loạt:")
                            st.dataframe(df)

                            csv_buffer = io.StringIO()
                            df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
                            csv = csv_buffer.getvalue().encode('utf-8')
                            st.download_button(
                                label="Tải báo cáo CSV",
                                data=csv,
                                file_name="batch_grading_report.csv",
                                mime="text/csv",
                                key="download_batch_report_csv_tab2"
                            )
                            st.success("Đã chấm xong tất cả bài và lưu kết quả trên Google Drive.")

                            st.subheader("Tải kết quả chi tiết cho sinh viên:")
                            response = service.files().list(q=f"'{graded_essays_folder_id}' in parents and trashed=false", spaces='drive').execute()
                            file_list = response.get('files', [])
                            if file_list:
                                # Tạo file ZIP chứa tất cả các file kết quả
                                zip_buffer = io.BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                    for file in file_list:
                                        if file['name'].endswith("_graded.docx") and "_graded_graded" not in file['name']:
                                            set_loading_cursor(True)
                                            with st.spinner(f"Đang xử lý file {file['name']}..."):
                                                file_content = download_file_from_drive(service, file['id'])
                                            set_loading_cursor(False)
                                            if file_content:
                                                zip_file.writestr(file['name'], file_content)

                                zip_buffer.seek(0)
                                st.download_button(
                                    label="Tải tất cả kết quả (ZIP)",
                                    data=zip_buffer,
                                    file_name="all_graded_essays.zip",
                                    mime="application/zip",
                                    key="download_all_graded_tab2"
                                )
                            else:
                                st.info("Chưa có kết quả chấm điểm nào được lưu.")
                        elif selected_essays and not st.session_state["start_grading"]:
                            st.info("Vui lòng nhấn 'Chấm bài' để bắt đầu chấm điểm.")
                        elif selected_essays:
                            st.info("Đang xử lý bài làm...")
                    else:
                        st.info("Chưa có bài làm nào được nộp.")

                # Sub-tab 3: Xem báo cáo
                with sub_tab3:
                    df = load_grading_report(service, reports_folder_id)
                    if df is not None:
                        st.subheader("Báo cáo điểm tổng hợp:")
                        st.dataframe(df)
                        csv_buffer = io.StringIO()
                        df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
                        csv = csv_buffer.getvalue().encode('utf-8')
                        st.download_button(
                            label="Tải báo cáo tổng hợp CSV",
                            data=csv,
                            file_name="grading_report_total.csv",
                            mime="text/csv",
                            key="download_total_report_csv_tab2"
                        )
                    else:
                        st.info("Chưa có báo cáo nào được lưu.")
    
    elif role == "student":
        # Lấy danh sách tất cả thư mục giảng viên
        response = service.files().list(q=f"'{root_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false", spaces='drive').execute()
        teacher_folders = response.get('files', [])
        
        # Lấy danh sách tất cả đề thi từ các giảng viên
        all_exams = []
        for teacher_folder in teacher_folders:
            username = teacher_folder['name'].replace("teacher_", "")
            exams_folder = find_file_in_folder(service, "exams", teacher_folder['id'])
            if exams_folder:
                exam_list = get_exam_list(service, exams_folder['id'])
                for exam in exam_list:
                    subject_code = exam.get("subject_code", "N/A")
                    term = exam.get("term", "N/A")
                    subject_name = exam.get("subject_name", "N/A")
                    display_name = f"{subject_code} - {term} - {subject_name} - {username}"
                    all_exams.append({
                        "display_name": display_name,
                        "exam_id": exam['exam_id'],
                        "secret_code": exam['secret_code']
                    })
        
        if all_exams:
            mssv = st.text_input("MSSV:", value=st.session_state["mssv"], key="mssv_input")
            full_name = st.text_input("Họ và Tên:", value=st.session_state["full_name"], key="full_name_input")
            st.session_state["mssv"] = mssv
            st.session_state["full_name"] = full_name
            
            if st.session_state["mssv"] and st.session_state["full_name"]:
                # Hiển thị danh sách đề thi với định dạng mới
                selected_exam = st.selectbox("Chọn đề thi:", [exam["display_name"] for exam in all_exams])
                secret_code = st.text_input("Nhập mã số bí mật:", type="password")
                
                if st.button("Xem đề thi"):
                    selected_exam_data = next(exam for exam in all_exams if exam["display_name"] == selected_exam)
                    if secret_code == selected_exam_data["secret_code"]:
                        st.session_state["selected_exam_id"] = selected_exam_data["exam_id"]
                        st.session_state["exam_access_granted"] = True
                        st.rerun()
                    else:
                        st.error("Mã số bí mật không đúng. Vui lòng thử lại.")
                
                if st.session_state.get("exam_access_granted", False):
                    tab1, tab2 = st.tabs(["Làm bài thi online", "Nộp bài"])
                    
                    with tab1:
                        if not st.session_state["start_exam"]:
                            if st.button("Làm bài"):
                                st.session_state["start_exam"] = True
                                st.session_state["current_num_questions"] = 1
                                st.rerun()
                        else:
                            st.subheader("Đề thi:")
                            file_id = st.session_state["selected_exam_id"]
                            viewer_url = f"https://drive.google.com/viewerng/viewer?embedded=true&url=https://drive.google.com/uc?id={file_id}"
                            pdf_display = f'<iframe src="{viewer_url}" width="100%" height="600px" frameborder="0"></iframe>'
                            st.markdown(pdf_display, unsafe_allow_html=True)
                            st.info("Nếu đề thi không hiển thị, vui lòng sử dụng nút 'Tải đề thi (PDF) nếu không xem được' để tải file về và xem.")
                            
                            set_loading_cursor(True)
                            with st.spinner("Đang tải đề thi..."):
                                exam_content = download_file_from_drive(service, file_id)
                            set_loading_cursor(False)
                            
                            st.download_button(
                                label="Tải đề thi (PDF) nếu không xem được",
                                data=exam_content,
                                file_name="de_thi.pdf",
                                mime="application/pdf"
                            )
                            
                            answers = []
                            for i in range(st.session_state["current_num_questions"]):
                                st.write(f"**Câu {i+1}**")
                                answer = st_quill(f"Câu {i+1}:", key=f"answer_{i}")
                                answers.append(answer)
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                if st.button("Thêm câu hỏi"):
                                    st.session_state["current_num_questions"] += 1
                                    st.rerun()
                            with col2:
                                if st.session_state["current_num_questions"] > 1:
                                    if st.button("Loại câu hỏi"):
                                        st.session_state["current_num_questions"] -= 1
                                        st.rerun()
                            
                            if st.button("Nộp bài"):
                                student_text = "\n".join([f"Câu {i+1}:\n{answer}" for i, answer in enumerate(answers) if answer])
                                filename = f"{st.session_state['mssv']}_{st.session_state['full_name']}.docx"
                                doc = docx.Document()
                                doc.add_paragraph(student_text)
                                doc_buffer = io.BytesIO()
                                doc.save(doc_buffer)
                                doc_buffer.seek(0)
                                
                                # Lưu bài làm vào thư mục essays của giảng viên tương ứng
                                teacher_username = selected_exam.split(" - ")[-1]
                                teacher_folder = get_or_create_folder(service, f"teacher_{teacher_username}", root_folder_id)
                                essays_folder = get_or_create_folder(service, "essays", teacher_folder)
                                upload_file_to_drive(service, doc_buffer.getvalue(), filename, essays_folder)
                                st.success(f"Bài làm đã được lưu trên Google Drive với tên: {filename}")
                                st.session_state["start_exam"] = False
                                st.session_state["current_num_questions"] = 1
                                st.session_state["exam_access_granted"] = False
                                st.rerun()
                    
                    with tab2:
                        uploaded_essay = st.file_uploader("Tải lên bài làm tự luận", type=["docx"])
                        if uploaded_essay:
                            filename = f"{st.session_state['mssv']}_{st.session_state['full_name']}.docx"
                            essay_content = uploaded_essay.read()
                            teacher_username = selected_exam.split(" - ")[-1]
                            teacher_folder = get_or_create_folder(service, f"teacher_{teacher_username}", root_folder_id)
                            essays_folder = get_or_create_folder(service, "essays", teacher_folder)
                            upload_file_to_drive(service, essay_content, filename, essays_folder)
                            st.success(f"Bài làm đã được lưu trên Google Drive với tên: {filename}")
        else:
            st.error("Không tìm thấy đề thi nào. Vui lòng liên hệ giáo viên.")
