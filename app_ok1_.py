#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Mar 24 07:48:40 2025

@author: phattai
"""

import streamlit as st
import requests
import docx
import os
import pandas as pd
from streamlit_quill import st_quill

# Sử dụng OpenRouter API miễn phí
API_URL = "https://openrouter.ai/api/v1/chat/completions"
API_KEY = "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"  # Thay bằng API key miễn phí từ OpenRouter

# Danh sách user giả lập
USERS = {
    "teacher": "1",
    "student": "1"
}


# Hàm kiểm tra đăng nhập
def login():
    st.session_state["logged_in"] = False
    user = st.text_input("Tên đăng nhập:")
    password = st.text_input("Mật khẩu:", type="password")
    if st.button("Đăng nhập"):
        if user in USERS and USERS[user] == password:
            st.session_state["logged_in"] = True
            st.session_state["user"] = user
            st.session_state["role"] = "teacher" if user == "teacher" else "student"
            st.success(f"Xin chào, {user}!")
        else:
            st.error("Sai tài khoản hoặc mật khẩu!")

# Hàm đăng xuất
def logout():
    st.session_state.clear()
    st.rerun()

# Hàm đọc file Word
def read_docx(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# Hàm lưu vào CSV
def save_to_csv(data):
    df = pd.DataFrame(data)
    if os.path.exists("grading_report.csv"):
        df.to_csv("grading_report.csv", mode="a", header=False, index=False)
    else:
        df.to_csv("grading_report.csv", mode="w", header=True, index=False)

# Hàm chấm điểm bài tự luận
def grade_essay(student_text, answer_text, student_name=None, mssv=None):
    prompt = f"""Bạn là giáo viên. Hãy chấm bài sau đây.
    \n\nĐáp án mẫu:\n{answer_text}
    \n\nBài làm của học sinh:\n{student_text}
    \n\nHãy đưa ra số điểm (thang 10) và nhận xét chi tiết. Định dạng điểm phải là: Điểm: [số điểm] (ví dụ: Điểm: 8.5)"""
    
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "mistralai/mistral-small-3.1-24b-instruct:free",
        "messages": [{"role": "system", "content": "Bạn là một giáo viên chấm bài chuyên nghiệp."},
                      {"role": "user", "content": prompt}],
        "temperature": 0.7
    }
    
    try:
        response = requests.post(API_URL, headers=headers, json=payload)
        if response.status_code == 200:
            grading_result = response.json()["choices"][0]["message"]["content"]
            if student_name and mssv:
                total_score = extract_score(grading_result)
                data = {
                    "MSSV": [mssv],
                    "Họ và Tên": [student_name],
                    "Điểm Tổng": [total_score],
                    "Kết quả chấm điểm": [grading_result]
                }
                save_to_csv(data)
            return grading_result
        else:
            st.error(f"Lỗi API: {response.status_code} - {response.json()}")
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"Lỗi kết nối mạng: {str(e)}")
        return None

# Hàm trích xuất điểm từ kết quả chấm
def extract_score(grading_result):
    import re
    match = re.search(r"Điểm:\s*(\d+(\.\d+)?)", grading_result, re.IGNORECASE)
    if match:
        return float(match.group(1))
    match = re.search(r"Score:\s*(\d+(\.\d+)?)", grading_result, re.IGNORECASE)
    if match:
        return float(match.group(1))
    match = re.search(r"(\d+(\.\d+)?)/10", grading_result)
    if match:
        return float(match.group(1))
    match = re.search(r"^\s*(\d+(\.\d+)?)\s*$", grading_result, re.MULTILINE)
    if match:
        return float(match.group(1))
    st.warning(f"Không thể trích xuất điểm từ kết quả: {grading_result}")
    return 0.0

# Hàm đọc báo cáo từ file CSV (sử dụng caching)
@st.cache_data
def load_grading_report():
    if os.path.exists("grading_report.csv"):
        return pd.read_csv("grading_report.csv")
    return None

# Cập nhật giao diện chính
if "logged_in" not in st.session_state or not st.session_state["logged_in"]:
    login()
else:
    st.write(f"Xin chào, {st.session_state['user']}!")
    if st.button("Đăng xuất"):
        logout()
    
    role = st.session_state.get("role", "student")
    
    if role == "teacher":
        os.makedirs("exams", exist_ok=True)

        # Phần tải đề thi và đáp án mẫu (đưa ra ngoài tab)
        st.subheader("Tải đề thi và đáp án")
        uploaded_exam = st.file_uploader("Tải lên đề thi tự luận", type=["docx"], key="exam")
        uploaded_answer = st.file_uploader("Tải lên đáp án mẫu", type=["docx"], key="answer")
        
        if uploaded_exam and uploaded_answer:
            exam_path = os.path.join("exams", "de_thi.docx")
            answer_path = os.path.join("exams", "dap_an.docx")
            
            with open(exam_path, "wb") as f:
                f.write(uploaded_exam.read())
            with open(answer_path, "wb") as f:
                f.write(uploaded_answer.read())
            
            st.success("Đề thi và đáp án đã được lưu.")

        # Tạo các tab cho các chức năng còn lại
        tab1, tab2, tab3 = st.tabs(["Chấm bài đơn", "Chấm bài hàng loạt", "Xem báo cáo"])

        # Tab 1: Chấm bài tự luận đơn
        with tab1:
            uploaded_essay = st.file_uploader("Tải lên bài làm tự luận của sinh viên", type=["docx"], key="single_essay")
            
            if uploaded_essay:
                answer_path = os.path.join("exams", "dap_an.docx")
                if os.path.exists(answer_path):
                    filename = uploaded_essay.name
                    try:
                        mssv, student_name = filename.replace(".docx", "").split("_", 1)
                    except ValueError:
                        st.error("Tên file không đúng định dạng 'MSSV_HọTên.docx'. Vui lòng kiểm tra lại.")
                    else:
                        student_text = read_docx(uploaded_essay)
                        answer_text = read_docx(answer_path)
                        result = grade_essay(student_text, answer_text, student_name, mssv)
                        
                        if result:
                            st.subheader("Kết quả chấm điểm:")
                            st.write(f"MSSV: {mssv}")
                            st.write(f"Họ và Tên: {student_name}")
                            st.write(result)
                            
                            os.makedirs("graded_essays", exist_ok=True)
                            graded_filename = f"{mssv}_{student_name}_graded.docx"
                            graded_path = os.path.join("graded_essays", graded_filename)
                            doc = docx.Document()
                            doc.add_paragraph(f"MSSV: {mssv}")
                            doc.add_paragraph(f"Họ và Tên: {student_name}")
                            doc.add_paragraph(result)
                            doc.save(graded_path)
                            
                            st.success(f"Kết quả đã được lưu vào file: {graded_filename}")
                            with open(graded_path, "rb") as f:
                                st.download_button(
                                    label="Tải kết quả chấm điểm",
                                    data=f,
                                    file_name=graded_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                else:
                    st.error("Không tìm thấy đáp án mẫu. Vui lòng tải lên đáp án trước.")

        # Tab 2: Chấm bài hàng loạt
        with tab2:
            if "uploaded_files" not in st.session_state:
                st.session_state["uploaded_files"] = []
            if "grading_results" not in st.session_state:
                st.session_state["grading_results"] = []

            uploaded_essays = st.file_uploader("Tải lên nhiều bài làm tự luận", type=["docx"], accept_multiple_files=True, key="batch_essays")
            
            current_files = [file.name for file in uploaded_essays] if uploaded_essays else []
            if current_files != st.session_state["uploaded_files"]:
                st.session_state["uploaded_files"] = current_files
                st.session_state["grading_results"] = []
                
                if uploaded_essays:
                    answer_path = os.path.join("exams", "dap_an.docx")
                    if os.path.exists(answer_path):
                        answer_text = read_docx(answer_path)
                        results = []
                        
                        # Tạo thư mục để lưu kết quả chấm điểm
                        os.makedirs("graded_essays", exist_ok=True)
                        
                        for idx, essay_file in enumerate(uploaded_essays, 1):
                            filename = essay_file.name
                            try:
                                mssv, student_name = filename.replace(".docx", "").split("_", 1)
                            except ValueError:
                                st.warning(f"Tên file {filename} không đúng định dạng 'MSSV_HọTên.docx'. Bỏ qua.")
                                continue
                            
                            student_text = read_docx(essay_file)
                            grading_result = grade_essay(student_text, answer_text, student_name, mssv)
                            
                            if grading_result:
                                total_score = extract_score(grading_result)
                                results.append({
                                    "STT": idx,
                                    "MSSV": mssv,
                                    "Họ và Tên": student_name,
                                    "Tổng điểm tự luận": total_score
                                })
                                
                                # Lưu kết quả chấm điểm vào file .docx
                                graded_filename = f"{mssv}_{student_name}_graded.docx"
                                graded_path = os.path.join("graded_essays", graded_filename)
                                doc = docx.Document()
                                doc.add_paragraph(f"MSSV: {mssv}")
                                doc.add_paragraph(f"Họ và Tên: {student_name}")
                                doc.add_paragraph(grading_result)
                                doc.save(graded_path)
                        
                        st.session_state["grading_results"] = results
                    else:
                        st.error("Không tìm thấy đáp án mẫu. Vui lòng tải lên đáp án trước.")

            # Hiển thị kết quả từ session state
            if st.session_state["grading_results"]:
                df = pd.DataFrame(st.session_state["grading_results"])
                st.subheader("Kết quả chấm điểm hàng loạt:")
                st.dataframe(df)
                
                csv = df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="Tải báo cáo CSV",
                    data=csv,
                    file_name="batch_grading_report.csv",
                    mime="text/csv"
                )
                st.success("Đã chấm xong tất cả bài và lưu vào file CSV.")

                # Hiển thị danh sách file kết quả chấm điểm để tải xuống
                st.subheader("Tải kết quả chi tiết cho sinh viên:")
                graded_files = [f for f in os.listdir("graded_essays") if f.endswith("_graded.docx")]
                if graded_files:
                    for graded_file in graded_files:
                        file_path = os.path.join("graded_essays", graded_file)
                        with open(file_path, "rb") as f:
                            st.download_button(
                                label=f"Tải kết quả: {graded_file}",
                                data=f,
                                file_name=graded_file,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                else:
                    st.info("Chưa có kết quả chấm điểm nào được lưu.")
            elif uploaded_essays:
                st.info("Đang xử lý bài làm...")

        # Tab 3: Xem báo cáo tổng hợp từ file CSV
        with tab3:
            df = load_grading_report()
            if df is not None:
                st.subheader("Báo cáo điểm tổng hợp:")
                st.dataframe(df)
                csv = df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="Tải báo cáo tổng hợp CSV",
                    data=csv,
                    file_name="grading_report_total.csv",
                    mime="text/csv"
                )
            else:
                st.info("Chưa có báo cáo nào được lưu.")
    
    elif role == "student":
        exam_path = os.path.join("exams", "de_thi.docx")
        if os.path.exists(exam_path):
            mssv = st.text_input("MSSV:")
            full_name = st.text_input("Họ và Tên:")
            if mssv and full_name:
                tab1, tab2 = st.tabs(["Làm bài thi online", "Nộp bài"])
                
                with tab1:
                    exam_text = read_docx(exam_path)
                    st.subheader("Đề thi:")
                    st.write(exam_text)
                
                    answers = []
                    for i, question in enumerate(exam_text.split("\n")):
                        st.write(f"**{question}**")
                        answer = st_quill(f"Câu {i+1}:", key=f"answer_{i}")
                        answers.append(answer)
                    
                    if st.button("Nộp bài"):
                        student_text = "\n".join(answers)
                        os.makedirs("essays", exist_ok=True)
                        filename = f"{mssv}_{full_name}.docx"
                        essay_path = os.path.join("essays", filename)
                        doc = docx.Document()
                        doc.add_paragraph(student_text)
                        doc.save(essay_path)
                        st.success(f"Bài làm đã được lưu với tên: {filename}")
                
                with tab2:
                    uploaded_essay = st.file_uploader("Tải lên bài làm tự luận", type=["docx"])
                    if uploaded_essay:
                        os.makedirs("essays", exist_ok=True)
                        filename = f"{mssv}_{full_name}.docx"
                        essay_path = os.path.join("essays", filename)
                        with open(essay_path, "wb") as f:
                            f.write(uploaded_essay.read())
                        st.success(f"Bài làm đã được lưu với tên: {filename}")